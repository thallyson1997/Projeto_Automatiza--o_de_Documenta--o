import os
import tempfile
import zipfile
import shutil
import io
from copy import deepcopy
from docx import Document
from docx.shared import Cm, Pt
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree


def gerar_documento_modelo2_base():
    """Retorna os bytes do arquivo modelo2.docx sem modificacoes."""
    modelo_path = os.path.join(os.path.dirname(__file__), '..', 'documento', 'modelo2.docx')

    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Arquivo modelo nao encontrado em {modelo_path}")

    with open(modelo_path, 'rb') as arquivo_modelo:
        return arquivo_modelo.read()


def gerar_documento_modelo2_empresa(empresa, data_inicio, data_fim, datas_formulario, imagens_formularios=None):
    """Gera documento com capa fixa e secao de formulario repetida para cada data."""
    if not datas_formulario:
        raise ValueError('E necessario informar ao menos uma data de formulario')

    if imagens_formularios is None:
        imagens_formularios = [{} for _ in datas_formulario]

    documento_expandido = _montar_estrutura_documento_modelo2(datas_formulario)
    return _aplicar_substituicoes_modelo2(
        documento_expandido,
        empresa,
        data_inicio,
        data_fim,
        imagens_formularios
    )


def _montar_estrutura_documento_modelo2(datas_formulario):
    """Monta o DOCX final antes das substituicoes: capa + (secao formulario * N)."""
    modelo_path = os.path.join(os.path.dirname(__file__), '..', 'documento', 'modelo2.docx')
    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Arquivo modelo nao encontrado em {modelo_path}")

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    temp_dir = tempfile.mkdtemp()

    try:
        base_dir = os.path.join(temp_dir, 'base')
        os.makedirs(base_dir, exist_ok=True)

        with zipfile.ZipFile(modelo_path, 'r') as zip_ref:
            zip_ref.extractall(base_dir)

        document_xml_path = os.path.join(base_dir, 'word', 'document.xml')
        tree = etree.parse(document_xml_path)
        root = tree.getroot()
        body = root.find('.//w:body', ns)

        elementos_sem_sectpr = [e for e in body if not e.tag.endswith('}sectPr')]
        idx_inicio_secao = _encontrar_inicio_secao_formulario(elementos_sem_sectpr)

        capa_elementos = elementos_sem_sectpr[:idx_inicio_secao]
        template_secao = elementos_sem_sectpr[idx_inicio_secao:]

        if not template_secao:
            raise ValueError('Secao de formulario vazia no modelo2.docx')

        # Limpa body mantendo apenas sectPr final
        for child in list(body):
            if not child.tag.endswith('}sectPr'):
                body.remove(child)

        # Insere capa fixa
        for elem in capa_elementos:
            _inserir_antes_do_sectpr(body, deepcopy(elem))

        # Insere secoes repetidas por formulario
        for i, data_formulario in enumerate(datas_formulario):
            for elem in template_secao:
                clone = deepcopy(elem)
                _replace_in_element_text(clone, '[DATA]', data_formulario)
                _inserir_antes_do_sectpr(body, clone)

        tree.write(document_xml_path, xml_declaration=True, encoding='UTF-8', standalone=True)

        resultado_path = os.path.join(temp_dir, 'resultado.docx')
        with zipfile.ZipFile(resultado_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for root_dir, dirs, files in os.walk(base_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, base_dir)
                    zip_ref.write(file_path, arcname)

        with open(resultado_path, 'rb') as arquivo:
            return arquivo.read()
    finally:
        try:
            shutil.rmtree(temp_dir)
        except:
            pass


def _aplicar_substituicoes_modelo2(documento_bytes, empresa, data_inicio, data_fim, imagens_formularios):
    """Aplica substituicoes de EMPRESA e PERIODO apos estruturar o documento final."""
    fd, tmp_doc_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)

    with open(tmp_doc_path, 'wb') as arquivo_tmp:
        arquivo_tmp.write(documento_bytes)

    try:
        doc = Document(tmp_doc_path)

        empresa_maiuscula = empresa.upper()
        primeira_ocorrencia_substituida = False
        idx_lanche = 0
        idx_ceia = 0
        idx_legenda_lanche = 0
        idx_legenda_ceia = 0
        idx_almoco = [0, 0, 0, 0]
        idx_jantar = [0, 0, 0, 0]
        idx_proteina_almoco = [0, 0, 0, 0]
        idx_proteina_jantar = [0, 0, 0, 0]
        idx_peso_almoco = [0, 0, 0, 0]
        idx_peso_jantar = [0, 0, 0, 0]
        idx_acompanhamento_almoco = [0, 0]
        idx_acompanhamento_jantar = [0, 0]

        def substituir_texto_paragrafo(paragrafo):
            nonlocal primeira_ocorrencia_substituida
            nonlocal idx_lanche, idx_ceia, idx_legenda_lanche, idx_legenda_ceia
            texto = paragrafo.text
            if (
                '[EMPRESA]' not in texto
                and '[DATA INICIO]' not in texto
                and '[DATA FIM]' not in texto
                and '[IMAGEM LANCHE]' not in texto
                and '[IMAGEM CEIA]' not in texto
                and '[LEGENDA LANCHE]' not in texto
                and '[LEGENDA CEIA]' not in texto
                and not any(f'[IMAGEM ALMOÇO {n}]' in texto for n in range(1, 5))
                and not any(f'[IMAGEM JANTAR {n}]' in texto for n in range(1, 5))
                and not any(f'[PROTEINA ALMOÇO {n}]' in texto for n in range(1, 5))
                and not any(f'[PROTEINA JANTAR {n}]' in texto for n in range(1, 5))
                and not any(f'[PESO ALMOÇO {n}]' in texto for n in range(1, 5))
                and not any(f'[PESO JANTAR {n}]' in texto for n in range(1, 5))
                and not any(f'[ACOMPANHAMENTO ALMOÇO {n}]' in texto for n in range(1, 3))
                and not any(f'[ACOMPANHAMENTO JANTAR {n}]' in texto for n in range(1, 3))
            ):
                return

            if '[EMPRESA]' in texto:
                if not primeira_ocorrencia_substituida:
                    texto = texto.replace('[EMPRESA]', empresa_maiuscula, 1)
                    primeira_ocorrencia_substituida = True
                texto = texto.replace('[EMPRESA]', empresa)

            texto = texto.replace('[DATA INICIO]', data_inicio)
            texto = texto.replace('[DATA FIM]', data_fim)

            if '[IMAGEM LANCHE]' in texto:
                dados = imagens_formularios[idx_lanche] if idx_lanche < len(imagens_formularios) else {}
                imagem_bytes = dados.get('imagem_lanche')
                idx_lanche += 1

                if imagem_bytes:
                    for run in paragrafo.runs:
                        run.text = ''
                    paragrafo.paragraph_format.space_before = Pt(0)
                    paragrafo.paragraph_format.space_after = Pt(0)
                    stream = io.BytesIO(imagem_bytes)
                    paragrafo.add_run().add_picture(stream, width=Cm(8), height=Cm(5))
                    return

                texto = texto.replace('[IMAGEM LANCHE]', '')

            if '[IMAGEM CEIA]' in texto:
                dados = imagens_formularios[idx_ceia] if idx_ceia < len(imagens_formularios) else {}
                imagem_bytes = dados.get('imagem_ceia')
                idx_ceia += 1

                if imagem_bytes:
                    for run in paragrafo.runs:
                        run.text = ''
                    paragrafo.paragraph_format.space_before = Pt(0)
                    paragrafo.paragraph_format.space_after = Pt(0)
                    stream = io.BytesIO(imagem_bytes)
                    paragrafo.add_run().add_picture(stream, width=Cm(8), height=Cm(5))
                    return

                texto = texto.replace('[IMAGEM CEIA]', '')

            if '[LEGENDA LANCHE]' in texto:
                dados = imagens_formularios[idx_legenda_lanche] if idx_legenda_lanche < len(imagens_formularios) else {}
                legenda = dados.get('legenda_lanche', '')
                idx_legenda_lanche += 1
                texto = texto.replace('[LEGENDA LANCHE]', legenda)

            if '[LEGENDA CEIA]' in texto:
                dados = imagens_formularios[idx_legenda_ceia] if idx_legenda_ceia < len(imagens_formularios) else {}
                legenda = dados.get('legenda_ceia', '')
                idx_legenda_ceia += 1
                texto = texto.replace('[LEGENDA CEIA]', legenda)

            for n in range(1, 5):
                placeholder_almoco = f'[IMAGEM ALMOÇO {n}]'
                if placeholder_almoco in texto:
                    dados = imagens_formularios[idx_almoco[n-1]] if idx_almoco[n-1] < len(imagens_formularios) else {}
                    imagem_bytes = dados.get(f'imagem_almoco_{n}')
                    idx_almoco[n-1] += 1
                    if imagem_bytes:
                        for run in paragrafo.runs:
                            run.text = ''
                        espaco = Pt(6) if n in (3, 4) else Pt(0)
                        paragrafo.paragraph_format.space_before = espaco
                        paragrafo.paragraph_format.space_after = espaco
                        stream = io.BytesIO(imagem_bytes)
                        paragrafo.add_run().add_picture(stream, width=Cm(8), height=Cm(5))
                        return
                    texto = texto.replace(placeholder_almoco, '')

            for n in range(1, 5):
                placeholder_jantar = f'[IMAGEM JANTAR {n}]'
                if placeholder_jantar in texto:
                    dados = imagens_formularios[idx_jantar[n-1]] if idx_jantar[n-1] < len(imagens_formularios) else {}
                    imagem_bytes = dados.get(f'imagem_jantar_{n}')
                    idx_jantar[n-1] += 1
                    if imagem_bytes:
                        for run in paragrafo.runs:
                            run.text = ''
                        espaco = Pt(6) if n in (3, 4) else Pt(0)
                        paragrafo.paragraph_format.space_before = espaco
                        paragrafo.paragraph_format.space_after = espaco
                        stream = io.BytesIO(imagem_bytes)
                        paragrafo.add_run().add_picture(stream, width=Cm(8), height=Cm(5))
                        return
                    texto = texto.replace(placeholder_jantar, '')

            for n in range(1, 5):
                placeholder_proteina_almoco = f'[PROTEINA ALMOÇO {n}]'
                if placeholder_proteina_almoco in texto:
                    dados = imagens_formularios[idx_proteina_almoco[n-1]] if idx_proteina_almoco[n-1] < len(imagens_formularios) else {}
                    proteina = dados.get(f'proteina_almoco_{n}', '')
                    idx_proteina_almoco[n-1] += 1
                    texto = texto.replace(placeholder_proteina_almoco, proteina)

            for n in range(1, 5):
                placeholder_proteina_jantar = f'[PROTEINA JANTAR {n}]'
                if placeholder_proteina_jantar in texto:
                    dados = imagens_formularios[idx_proteina_jantar[n-1]] if idx_proteina_jantar[n-1] < len(imagens_formularios) else {}
                    proteina = dados.get(f'proteina_jantar_{n}', '')
                    idx_proteina_jantar[n-1] += 1
                    texto = texto.replace(placeholder_proteina_jantar, proteina)

            for n in range(1, 5):
                placeholder_peso_almoco = f'[PESO ALMOÇO {n}]'
                if placeholder_peso_almoco in texto:
                    dados = imagens_formularios[idx_peso_almoco[n-1]] if idx_peso_almoco[n-1] < len(imagens_formularios) else {}
                    peso = dados.get(f'peso_almoco_{n}', '')
                    idx_peso_almoco[n-1] += 1
                    texto = texto.replace(placeholder_peso_almoco, peso)

            for n in range(1, 5):
                placeholder_peso_jantar = f'[PESO JANTAR {n}]'
                if placeholder_peso_jantar in texto:
                    dados = imagens_formularios[idx_peso_jantar[n-1]] if idx_peso_jantar[n-1] < len(imagens_formularios) else {}
                    peso = dados.get(f'peso_jantar_{n}', '')
                    idx_peso_jantar[n-1] += 1
                    texto = texto.replace(placeholder_peso_jantar, peso)

            for n in range(1, 3):
                placeholder_acompanhamento_almoco = f'[ACOMPANHAMENTO ALMOÇO {n}]'
                if placeholder_acompanhamento_almoco in texto:
                    dados = imagens_formularios[idx_acompanhamento_almoco[n-1]] if idx_acompanhamento_almoco[n-1] < len(imagens_formularios) else {}
                    acompanhamento = dados.get(f'acompanhamento_almoco_{n}', '')
                    idx_acompanhamento_almoco[n-1] += 1
                    texto = texto.replace(placeholder_acompanhamento_almoco, acompanhamento)

            for n in range(1, 3):
                placeholder_acompanhamento_jantar = f'[ACOMPANHAMENTO JANTAR {n}]'
                if placeholder_acompanhamento_jantar in texto:
                    dados = imagens_formularios[idx_acompanhamento_jantar[n-1]] if idx_acompanhamento_jantar[n-1] < len(imagens_formularios) else {}
                    acompanhamento = dados.get(f'acompanhamento_jantar_{n}', '')
                    idx_acompanhamento_jantar[n-1] += 1
                    texto = texto.replace(placeholder_acompanhamento_jantar, acompanhamento)

            for run in paragrafo.runs:
                run.text = ''

            if paragrafo.runs:
                paragrafo.runs[0].text = texto
            else:
                paragrafo.add_run(texto)

        def iterar_blocos(parent):
            if isinstance(parent, DocxDocument):
                parent_elemento = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elemento = parent._tc
            else:
                return

            for child in parent_elemento.iterchildren():
                if child.tag.endswith('}p'):
                    yield Paragraph(child, parent)
                elif child.tag.endswith('}tbl'):
                    yield Table(child, parent)

        def processar_tabela(tabela):
            for linha in tabela.rows:
                for celula in linha.cells:
                    for bloco in iterar_blocos(celula):
                        if isinstance(bloco, Paragraph):
                            substituir_texto_paragrafo(bloco)
                        elif isinstance(bloco, Table):
                            processar_tabela(bloco)

        for bloco in iterar_blocos(doc):
            if isinstance(bloco, Paragraph):
                substituir_texto_paragrafo(bloco)
            elif isinstance(bloco, Table):
                processar_tabela(bloco)

        doc.save(tmp_doc_path)
        with open(tmp_doc_path, 'rb') as arquivo:
            return arquivo.read()
    finally:
        try:
            os.unlink(tmp_doc_path)
        except:
            pass

def _encontrar_inicio_secao_formulario(elementos_sem_sectpr):
    """Encontra onde a secao repetivel (paginas 2 e 3) comeca no modelo."""
    for i, elem in enumerate(elementos_sem_sectpr):
        texto = ''.join(elem.itertext())
        if '[DATA]' in texto:
            return max(0, i - 1)

    raise ValueError(
        'Nao foi possivel encontrar o inicio da secao de formulario no modelo2.docx. '
        'Verifique se o placeholder [DATA] existe nas paginas de formulario.'
    )


def _replace_in_element_text(elemento, old, new):
    """Substitui texto em todos os nos textuais do elemento XML."""
    for node in elemento.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        if node.text:
            node.text = node.text.replace(old, new)


def _criar_paragrafo_quebra_pagina():
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p = etree.Element(f'{{{ns_w}}}p')
    r = etree.SubElement(p, f'{{{ns_w}}}r')
    br = etree.SubElement(r, f'{{{ns_w}}}br')
    br.set(f'{{{ns_w}}}type', 'page')
    return p


def _inserir_antes_do_sectpr(body, elemento):
    for idx, child in enumerate(body):
        if child.tag.endswith('}sectPr'):
            body.insert(idx, elemento)
            return
    body.append(elemento)
