from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import tempfile
import shutil


def gerar_documento(unidade, data, legenda, imagens=None):
    """
    Carrega o modelo.docx, substitui os placeholders e insere as imagens
    e retorna o documento modificado com ajuste de altura.
    
    Args:
        unidade (str): Texto para substituir [UNIDADE]
        data (str): Data formatada (DD.MM.YYYY) para substituir [DATA]
        legenda (str): Texto para substituir [LEGENDA]
        imagens (list): Lista com até 4 imagens em bytes
    
    Returns:
        bytes: Documento Word em bytes
    """
    try:
        # Caminho do modelo
        modelo_path = os.path.join(os.path.dirname(__file__), '..', 'documento', 'modelo.docx')
        
        # Verifica se o arquivo modelo existe
        if not os.path.exists(modelo_path):
            raise FileNotFoundError(f"Arquivo modelo não encontrado em {modelo_path}")
        
        if imagens is None:
            imagens = []
        
        # Abre o documento modelo
        doc = Document(modelo_path)
        
        # Substitui placeholders em parágrafos top-level
        for paragrafo in doc.paragraphs:
            texto_completo = paragrafo.text
            
            if '[UNIDADE]' in texto_completo or '[DATA]' in texto_completo or '[LEGENDA]' in texto_completo:
                texto_completo = texto_completo.replace('[UNIDADE]', unidade)
                texto_completo = texto_completo.replace('[DATA]', data)
                texto_completo = texto_completo.replace('[LEGENDA]', legenda)
                
                for run in paragrafo.runs:
                    run.text = ''
                
                if paragrafo.runs:
                    paragrafo.runs[0].text = texto_completo
                else:
                    paragrafo.add_run(texto_completo)
        
        # Procura e substitui em tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        texto_completo = paragrafo.text
                        
                        # Se encontrar [IMAGENS], insere as imagens nessa célula
                        if '[IMAGENS]' in texto_completo:
                            # Limpa o parágrafo
                            for run in paragrafo.runs:
                                run.text = ''
                            
                            # Insere as imagens com altura de 7.5 cm
                            if imagens:
                                inserir_imagens_na_celula(celula, imagens, 7.5)
                        
                        # Substitui outros placeholders
                        elif '[UNIDADE]' in texto_completo or '[DATA]' in texto_completo or '[LEGENDA]' in texto_completo:
                            texto_completo = texto_completo.replace('[UNIDADE]', unidade)
                            texto_completo = texto_completo.replace('[DATA]', data)
                            texto_completo = texto_completo.replace('[LEGENDA]', legenda)
                            
                            for run in paragrafo.runs:
                                run.text = ''
                            
                            if paragrafo.runs:
                                paragrafo.runs[0].text = texto_completo
                            else:
                                paragrafo.add_run(texto_completo)
        
        # Salva o documento em um arquivo temporário e depois lê em bytes
        fd, tmp_doc_path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        
        try:
            doc.save(tmp_doc_path)
            with open(tmp_doc_path, 'rb') as f:
                doc_bytes_content = f.read()
            return doc_bytes_content
        finally:
            try:
                os.unlink(tmp_doc_path)
            except:
                pass
    
    except Exception as e:
        raise Exception(f"Erro ao gerar documento: {str(e)}")


def inserir_imagens_na_celula(celula, imagens, altura_imagem_cm=7.5):
    """
    Insere as imagens em uma célula de tabela de forma apropriada:
    - 1 imagem: centralizada
    - 2 imagens: uma em cima da outra
    - 3 imagens: 2 em cima, 1 embaixo
    - 4 imagens: 2 em cima, 2 embaixo
    """
    num_imagens = len(imagens)
    if num_imagens == 0:
        return
    
    # Remove todos os runs do primeiro parágrafo (limpeza completa)
    primeiro_paragrafo = celula.paragraphs[0]
    for run in list(primeiro_paragrafo.runs):
        run_elem = run._element
        run_elem.getparent().remove(run_elem)
    
    primeiro_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remover_espacamento_paragrafo(primeiro_paragrafo)
    centralizar_celula_verticalmente(celula)
    
    # Define dimensões baseado na quantidade de imagens
    if num_imagens <= 2:
        largura_cm = 9.6
    else:
        largura_cm = 6.8
    
    largura = Inches(largura_cm / 2.54)
    altura = Inches(altura_imagem_cm / 2.54)
    
    if num_imagens == 1:
        _adicionar_imagem_temp(primeiro_paragrafo, imagens[0], largura, altura)
    
    elif num_imagens == 2:
        _adicionar_imagem_temp(primeiro_paragrafo, imagens[0], largura, altura)
        primeiro_paragrafo.paragraph_format.space_after = Pt(6)
        
        novo_p = celula.add_paragraph()
        novo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        remover_espacamento_paragrafo(novo_p)
        _adicionar_imagem_temp(novo_p, imagens[1], largura, altura)
    
    elif num_imagens == 3:
        _adicionar_imagem_temp(primeiro_paragrafo, imagens[0], largura, altura)
        primeiro_paragrafo.add_run('  ')
        _adicionar_imagem_temp(primeiro_paragrafo, imagens[1], largura, altura)
        primeiro_paragrafo.paragraph_format.space_after = Pt(6)
        
        novo_p = celula.add_paragraph()
        novo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        remover_espacamento_paragrafo(novo_p)
        _adicionar_imagem_temp(novo_p, imagens[2], largura, altura)
    
    elif num_imagens >= 4:
        _adicionar_imagem_temp(primeiro_paragrafo, imagens[0], largura, altura)
        primeiro_paragrafo.add_run('  ')
        _adicionar_imagem_temp(primeiro_paragrafo, imagens[1], largura, altura)
        primeiro_paragrafo.paragraph_format.space_after = Pt(6)
        
        novo_p = celula.add_paragraph()
        novo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        remover_espacamento_paragrafo(novo_p)
        _adicionar_imagem_temp(novo_p, imagens[2], largura, altura)
        novo_p.add_run('  ')
        _adicionar_imagem_temp(novo_p, imagens[3], largura, altura)


def _adicionar_imagem_temp(paragrafo, imagem_bytes, largura, altura):
    """Helper: Adiciona imagem usando arquivo temporário"""
    fd, tmp_path = tempfile.mkstemp(suffix='.png')
    os.close(fd)
    
    try:
        with open(tmp_path, 'wb') as f:
            f.write(imagem_bytes)
        paragrafo.add_run().add_picture(tmp_path, width=largura, height=altura)
    finally:
        try:
            os.unlink(tmp_path)
        except:
            pass


def obter_largura_celula(celula):
    """
    Obtém a largura da célula em polegadas.
    Se não conseguir obter, retorna um valor padrão.
    """
    try:
        tcPr = celula._element.tcPr
        if tcPr is not None:
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                # Obtém o valor em twips (1/20 de um ponto)
                w_val = tcW.get(qn('w:w'))
                if w_val:
                    # Converte twips para polegadas (1440 twips = 1 polegada)
                    twips = int(w_val)
                    polegadas = twips / 1440
                    return polegadas
    except:
        pass
    
    # Retorna valor padrão se não conseguir obter
    return 3.0  # 3 polegadas como padrão


def centralizar_celula_verticalmente(celula):
    """Centraliza o conteúdo da célula verticalmente"""
    tcPr = celula._element.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        celula._element.insert(0, tcPr)
    
    # Remove vAlign existente se houver
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is not None:
        tcPr.remove(vAlign)
    
    # Adiciona vAlign com valor center
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def remover_espacamento_paragrafo(paragrafo):
    """Remove espaçamento antes e depois do parágrafo"""
    pPr = paragrafo._element.get_or_add_pPr()
    pSpacing = pPr.find(qn('w:spacing'))
    
    if pSpacing is None:
        pSpacing = OxmlElement('w:spacing')
        pPr.append(pSpacing)
    
    # Define espaçamento antes e depois como 0
    pSpacing.set(qn('w:before'), '0')
    pSpacing.set(qn('w:after'), '0')
    pSpacing.set(qn('w:line'), '240')  # Single spacing (240 twips)
    pSpacing.set(qn('w:lineRule'), 'auto')


def adicionar_imagem_ao_paragrafo(paragrafo, imagem_bytes, largura_inches):
    """Adiciona uma imagem a um parágrafo"""
    imagem_stream = io.BytesIO(imagem_bytes)
    paragrafo.add_run().add_picture(imagem_stream, width=Inches(largura_inches))


def gerar_documento_multiplo(formularios):
    """
    Gera um único documento com múltiplas páginas, uma para cada formulário.
    Mantém a mesma estrutura e formatação do documento original.
    
    Args:
        formularios (list): Lista de dicts com chaves unidade, data, legenda, imagens
    
    Returns:
        bytes: Documento Word em bytes com múltiplas páginas
    """
    try:
        import zipfile
        from lxml import etree
        
        modelo_path = os.path.join(os.path.dirname(__file__), '..', 'documento', 'modelo.docx')
        
        if not os.path.exists(modelo_path):
            raise FileNotFoundError(f"Arquivo modelo não encontrado em {modelo_path}")
        
        # Cria documentos temporários para cada formulário
        temp_docs = []
        
        for form_data in formularios:
            # Gera documento individual usando gerar_documento
            doc_bytes = gerar_documento(
                form_data['unidade'],
                form_data['data'],
                form_data['legenda'],
                form_data['imagens']
            )
            
            # Salva em arquivo temporário
            fd, tmp_path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            
            with open(tmp_path, 'wb') as f:
                f.write(doc_bytes)
            
            temp_docs.append(tmp_path)
        
        # Agora merges todos os documentos usando o primeiro como base
        # e adicionando conteúdo dos demais
        final_temp_dir = tempfile.mkdtemp()
        final_zip_path = os.path.join(final_temp_dir, 'merged.docx')
        
        try:
            # Extrai o primeiro documento como base
            base_doc_path = temp_docs[0]
            
            with zipfile.ZipFile(base_doc_path, 'r') as zip_ref:
                zip_ref.extractall(final_temp_dir)
            
            # Lê o document.xml e o document.xml.rels do documento base
            doc_xml_path = os.path.join(final_temp_dir, 'word', 'document.xml')
            rels_path = os.path.join(final_temp_dir, 'word', '_rels', 'document.xml.rels')
            
            doc_tree = etree.parse(doc_xml_path)
            doc_root = doc_tree.getroot()
            
            rels_tree = etree.parse(rels_path)
            rels_root = rels_tree.getroot()
            
            # Namespace do documento
            ns_doc = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Encontra o body do documento
            body = doc_root.find('.//w:body', ns_doc)
            
            # Processa os documentos adicionais (índice 1 em diante)
            media_counter = 1
            
            # Conta quantas imagens já existem
            media_files = os.listdir(os.path.join(final_temp_dir, 'word', 'media'))
            if media_files:
                media_counter = max([int(f.replace('image', '').replace('.png', '')) for f in media_files if f.startswith('image')]) + 1
            
            for extra_doc_path in temp_docs[1:]:
                # Extrai o documento adicional em diretório temporário
                extra_temp_dir = tempfile.mkdtemp()
                
                with zipfile.ZipFile(extra_doc_path, 'r') as zip_ref:
                    zip_ref.extractall(extra_temp_dir)
                
                # Lê o document.xml do documento adicional
                extra_doc_xml_path = os.path.join(extra_temp_dir, 'word', 'document.xml')
                extra_doc_tree = etree.parse(extra_doc_xml_path)
                extra_doc_root = extra_doc_tree.getroot()
                extra_body = extra_doc_root.find('.//w:body', ns_doc)
                
                # Lê as relações do documento adicional
                extra_rels_path = os.path.join(extra_temp_dir, 'word', '_rels', 'document.xml.rels')
                extra_rels_tree = etree.parse(extra_rels_path)
                extra_rels_root = extra_rels_tree.getroot()
                
                # Mapeamento de IDs antigos para novos (para relações)
                rel_id_mapping = {}
                ns_rels = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                
                # Copia as relações de imagem do documento adicional
                for extra_rel in extra_rels_root.findall('.//r:Relationship[@Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"]', ns_rels):
                    old_rel_id = extra_rel.get('Id')
                    old_target = extra_rel.get('Target')
                    
                    # Cria novo ID de relação
                    new_rel_id = f'rId{999 + media_counter}'
                    rel_id_mapping[old_rel_id] = new_rel_id
                    
                    # Copia o arquivo de imagem
                    old_image_path = os.path.join(extra_temp_dir, 'word', old_target)
                    new_image_name = f'image{media_counter}.png'
                    new_image_path = os.path.join(final_temp_dir, 'word', 'media', new_image_name)
                    
                    if os.path.exists(old_image_path):
                        shutil.copy(old_image_path, new_image_path)
                    
                    # Adiciona a relação ao documento final
                    new_rel = etree.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
                    new_rel.set('Id', new_rel_id)
                    new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')
                    new_rel.set('Target', f'media/{new_image_name}')
                    rels_root.append(new_rel)
                    
                    media_counter += 1
                
                # Adiciona quebra de página antes de copiar conteúdo adicional
                page_break = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
                pPr = etree.SubElement(page_break, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                br = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
                br.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'page')
                body.append(page_break)
                
                # Copia os elementos do body do documento adicional
                for elemento in extra_body:
                    # Atualiza as referências de relação nos elementos copiados
                    _atualizar_refs_nos_elementos(elemento, rel_id_mapping)
                    body.append(elemento)
                
                # Limpa o diretório temporário
                shutil.rmtree(extra_temp_dir)
            
            # Salva o XML atualizado
            doc_tree.write(doc_xml_path, xml_declaration=True, encoding='UTF-8', standalone=True)
            rels_tree.write(rels_path, xml_declaration=True, encoding='UTF-8', standalone=True)
            
            # Recria o arquivo DOCX
            with zipfile.ZipFile(final_zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                for root, dirs, files in os.walk(final_temp_dir):
                    for file in files:
                        if file == 'merged.docx':
                            continue
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, final_temp_dir)
                        zip_ref.write(file_path, arcname)
            
            # Lê o arquivo final
            with open(final_zip_path, 'rb') as f:
                result_bytes = f.read()
            
            return result_bytes
        
        finally:
            # Limpa os arquivos temporários
            for doc_path in temp_docs:
                try:
                    os.unlink(doc_path)
                except:
                    pass
            
            try:
                shutil.rmtree(final_temp_dir)
            except:
                pass
    
    except Exception as e:
        raise Exception(f"Erro ao gerar documento múltiplo: {str(e)}")


def _atualizar_refs_nos_elementos(elemento, rel_id_mapping):
    """Atualiza as referências de relação nos elementos XML copiados"""
    ns_r = {'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
    
    # Procura por atributos r:embed ou r:link
    for key in elemento.attrib:
        if 'embed' in key or 'link' in key:
            old_id = elemento.get(key)
            if old_id in rel_id_mapping:
                elemento.set(key, rel_id_mapping[old_id])
    
    # Recursivamente processa elementos filhos
    for child in elemento:
        _atualizar_refs_nos_elementos(child, rel_id_mapping)

