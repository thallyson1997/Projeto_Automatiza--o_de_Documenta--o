import os
import tempfile
import zipfile
import shutil
import io
from copy import deepcopy
from docx import Document
from docx.shared import Cm, Pt
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree
from PIL import Image


def gerar_documento_modelo3_alipen(datas_formulario, unidades_formulario=None, imagens_formularios=None):
    """Gera documento sem capa, apenas com secoes de formulario repetidas (modelo ALIPEN)."""
    if not datas_formulario:
        raise ValueError('E necessario informar ao menos uma data de formulario')

    if unidades_formulario is None:
        unidades_formulario = ['' for _ in datas_formulario]
    
    if imagens_formularios is None:
        imagens_formularios = [{} for _ in datas_formulario]

    documento_expandido = _montar_estrutura_documento_modelo3(datas_formulario, unidades_formulario)
    return _aplicar_substituicoes_modelo3(documento_expandido, imagens_formularios)


def _montar_estrutura_documento_modelo3(datas_formulario, unidades_formulario):
    """Monta DOCX sem capa: apenas secoes de formulario repetidas."""
    modelo_path = os.path.join(os.path.dirname(__file__), '..', 'documento', 'modelo3.docx')
    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Arquivo modelo nao encontrado em {modelo_path}")

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    temp_dir = tempfile.mkdtemp()

    try:
        base_dir = os.path.join(temp_dir, 'base')
        os.makedirs(base_dir, exist_ok=True)

        with zipfile.ZipFile(modelo_path, 'r') as zip_ref:
            zip_ref.extractall(base_dir)

        document_xml_path = os.path.join(base_dir, 'word', 'document.xml')
        tree = etree.parse(document_xml_path)
        root = tree.getroot()
        body = root.find('.//w:body', ns)

        elementos_sem_sectpr = [e for e in body if not e.tag.endswith('}sectPr')]
        idx_inicio_secao = _encontrar_inicio_secao_formulario(elementos_sem_sectpr)

        template_secao = elementos_sem_sectpr[idx_inicio_secao:]

        if not template_secao:
            raise ValueError('Secao de formulario vazia no modelo2.docx')

        # Limpa body mantendo apenas sectPr final
        for child in list(body):
            if not child.tag.endswith('}sectPr'):
                body.remove(child)

        # Insere secoes repetidas por formulario (SEM CAPA)
        for i, (data_formulario, unidade_formulario) in enumerate(zip(datas_formulario, unidades_formulario)):
            if i > 0:
                _inserir_antes_do_sectpr(body, _criar_paragrafo_quebra_pagina())

            for elem in template_secao:
                clone = deepcopy(elem)
                _replace_in_element_text(clone, '[DATA]', data_formulario)
                _replace_in_element_text(clone, '[UNIDADE]', unidade_formulario)
                _inserir_antes_do_sectpr(body, clone)

        tree.write(document_xml_path, xml_declaration=True, encoding='UTF-8', standalone=True)

        resultado_path = os.path.join(temp_dir, 'resultado.docx')
        with zipfile.ZipFile(resultado_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for root_dir, dirs, files in os.walk(base_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, base_dir)
                    zip_ref.write(file_path, arcname)

        with open(resultado_path, 'rb') as arquivo:
            return arquivo.read()
    finally:
        try:
            shutil.rmtree(temp_dir)
        except:
            pass


def _calcular_dimensoes_imagem(imagem_bytes):
    """
    Calcula dimensões proporcionais da imagem baseado em orientação.
    
    Paisagem (largura > altura): largura = 8cm, altura = máx 6cm (proporcional)
    Retrato (altura > largura): altura = 6cm, largura = proporcional
    
    Retorna: (largura, altura, é_paisagem)
    """
    try:
        stream = io.BytesIO(imagem_bytes)
        img = Image.open(stream)
        largura_original, altura_original = img.size
        
        if largura_original == 0 or altura_original == 0:
            return Cm(8), Cm(6), True
        
        proporcao = altura_original / largura_original
        
        # Paisagem: largura > altura
        if largura_original > altura_original:
            largura = Cm(8)
            altura_calc = 8 * proporcao
            altura = Cm(min(altura_calc, 6))
            return largura, altura, True
        # Retrato: altura > largura
        else:
            altura = Cm(6)
            largura_calc = 6 / proporcao
            largura = Cm(largura_calc)
            return largura, altura, False
    except Exception:
        # Se houver erro ao detectar dimensões, usa padrão
        return Cm(8), Cm(6), True


def _aplicar_substituicoes_modelo3(documento_bytes, imagens_formularios):
    """Aplica substituicoes para modelo3 ALIPEN: café, lanche, almoço, jantar"""
    fd, tmp_doc_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)

    with open(tmp_doc_path, 'wb') as arquivo_tmp:
        arquivo_tmp.write(documento_bytes)

    try:
        doc = Document(tmp_doc_path)

        idx_cafe = 0
        idx_legenda_cafe = 0
        idx_lanche = 0
        idx_legenda_lanche = 0
        idx_almoco = 0
        idx_jantar = 0

        def substituir_texto_paragrafo(paragrafo):
            nonlocal idx_cafe, idx_legenda_cafe, idx_lanche, idx_legenda_lanche
            nonlocal idx_almoco, idx_jantar
            
            texto = paragrafo.text
            if (
                '[IMAGEM CAFÉ]' not in texto
                and '[IMAGEM LANCHE]' not in texto
                and '[IMAGEM ALMOÇO]' not in texto
                and '[IMAGEM JANTAR]' not in texto
                and '[LEGENDA CAFÉ]' not in texto
                and '[LEGENDA LANCHE]' not in texto
                and '[PROTEINA ALMOÇO]' not in texto
                and '[PROTEINA JANTAR]' not in texto
                and '[PESO ALMOÇO]' not in texto
                and '[PESO JANTAR]' not in texto
                and '[ACOMPANHAMENTO ALMOÇO]' not in texto
                and '[ACOMPANHAMENTO JANTAR]' not in texto
            ):
                return

            if '[IMAGEM CAFÉ]' in texto:
                dados = imagens_formularios[idx_cafe] if idx_cafe < len(imagens_formularios) else {}
                imagem_bytes = dados.get('imagem_cafe')
                idx_cafe += 1
                if imagem_bytes:
                    for run in paragrafo.runs:
                        run.text = ''
                    stream = io.BytesIO(imagem_bytes)
                    largura, altura, e_paisagem = _calcular_dimensoes_imagem(imagem_bytes)
                    if e_paisagem:
                        paragrafo.paragraph_format.space_before = Pt(14)
                    paragrafo.add_run().add_picture(stream, width=largura, height=altura)
                    return
                texto = texto.replace('[IMAGEM CAFÉ]', '')

            if '[IMAGEM LANCHE]' in texto:
                dados = imagens_formularios[idx_lanche] if idx_lanche < len(imagens_formularios) else {}
                imagem_bytes = dados.get('imagem_lanche')
                idx_lanche += 1
                if imagem_bytes:
                    for run in paragrafo.runs:
                        run.text = ''
                    stream = io.BytesIO(imagem_bytes)
                    largura, altura, e_paisagem = _calcular_dimensoes_imagem(imagem_bytes)
                    if e_paisagem:
                        paragrafo.paragraph_format.space_before = Pt(14)
                    paragrafo.add_run().add_picture(stream, width=largura, height=altura)
                    return
                texto = texto.replace('[IMAGEM LANCHE]', '')

            if '[LEGENDA CAFÉ]' in texto:
                dados = imagens_formularios[idx_legenda_cafe] if idx_legenda_cafe < len(imagens_formularios) else {}
                legenda = dados.get('legenda_cafe', '')
                idx_legenda_cafe += 1
                texto = texto.replace('[LEGENDA CAFÉ]', legenda)

            if '[LEGENDA LANCHE]' in texto:
                dados = imagens_formularios[idx_legenda_lanche] if idx_legenda_lanche < len(imagens_formularios) else {}
                legenda = dados.get('legenda_lanche', '')
                idx_legenda_lanche += 1
                texto = texto.replace('[LEGENDA LANCHE]', legenda)

            if '[IMAGEM ALMOÇO]' in texto:
                dados = imagens_formularios[idx_almoco] if idx_almoco < len(imagens_formularios) else {}
                imagem_bytes = dados.get('imagem_almoco')
                idx_almoco += 1
                if imagem_bytes:
                    for run in paragrafo.runs:
                        run.text = ''
                    stream = io.BytesIO(imagem_bytes)
                    largura, altura, e_paisagem = _calcular_dimensoes_imagem(imagem_bytes)
                    if e_paisagem:
                        paragrafo.paragraph_format.space_before = Pt(14)
                    paragrafo.add_run().add_picture(stream, width=largura, height=altura)
                    return
                texto = texto.replace('[IMAGEM ALMOÇO]', '')

            if '[IMAGEM JANTAR]' in texto:
                dados = imagens_formularios[idx_jantar] if idx_jantar < len(imagens_formularios) else {}
                imagem_bytes = dados.get('imagem_jantar')
                idx_jantar += 1
                if imagem_bytes:
                    for run in paragrafo.runs:
                        run.text = ''
                    stream = io.BytesIO(imagem_bytes)
                    largura, altura, e_paisagem = _calcular_dimensoes_imagem(imagem_bytes)
                    if e_paisagem:
                        paragrafo.paragraph_format.space_before = Pt(14)
                    paragrafo.add_run().add_picture(stream, width=largura, height=altura)
                    return
                texto = texto.replace('[IMAGEM JANTAR]', '')

            if '[PROTEINA ALMOÇO]' in texto:
                dados = imagens_formularios[idx_almoco] if idx_almoco < len(imagens_formularios) else {}
                proteina = dados.get('proteina_almoco', '')
                texto = texto.replace('[PROTEINA ALMOÇO]', proteina)

            if '[PROTEINA JANTAR]' in texto:
                dados = imagens_formularios[idx_jantar] if idx_jantar < len(imagens_formularios) else {}
                proteina = dados.get('proteina_jantar', '')
                texto = texto.replace('[PROTEINA JANTAR]', proteina)

            if '[PESO ALMOÇO]' in texto:
                dados = imagens_formularios[idx_almoco] if idx_almoco < len(imagens_formularios) else {}
                peso = dados.get('peso_almoco', '')
                texto = texto.replace('[PESO ALMOÇO]', peso)

            if '[PESO JANTAR]' in texto:
                dados = imagens_formularios[idx_jantar] if idx_jantar < len(imagens_formularios) else {}
                peso = dados.get('peso_jantar', '')
                texto = texto.replace('[PESO JANTAR]', peso)

            if '[ACOMPANHAMENTO ALMOÇO]' in texto:
                dados = imagens_formularios[idx_almoco] if idx_almoco < len(imagens_formularios) else {}
                acompanhamento = dados.get('acompanhamento_almoco', '')
                texto = texto.replace('[ACOMPANHAMENTO ALMOÇO]', acompanhamento)

            if '[ACOMPANHAMENTO JANTAR]' in texto:
                dados = imagens_formularios[idx_jantar] if idx_jantar < len(imagens_formularios) else {}
                acompanhamento = dados.get('acompanhamento_jantar', '')
                texto = texto.replace('[ACOMPANHAMENTO JANTAR]', acompanhamento)

            for run in paragrafo.runs:
                run.text = ''

            if paragrafo.runs:
                paragrafo.runs[0].text = texto
            else:
                paragrafo.add_run(texto)

        def iterar_blocos(parent):
            if isinstance(parent, DocxDocument):
                parent_elemento = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elemento = parent._tc
            else:
                return

            for child in parent_elemento.iterchildren():
                if child.tag.endswith('}p'):
                    yield Paragraph(child, parent)
                elif child.tag.endswith('}tbl'):
                    yield Table(child, parent)

        def processar_tabela(tabela):
            for linha in tabela.rows:
                for celula in linha.cells:
                    for bloco in iterar_blocos(celula):
                        if isinstance(bloco, Paragraph):
                            substituir_texto_paragrafo(bloco)
                        elif isinstance(bloco, Table):
                            processar_tabela(bloco)

        for bloco in iterar_blocos(doc):
            if isinstance(bloco, Paragraph):
                substituir_texto_paragrafo(bloco)
            elif isinstance(bloco, Table):
                processar_tabela(bloco)

        doc.save(tmp_doc_path)
        with open(tmp_doc_path, 'rb') as arquivo:
            return arquivo.read()
    finally:
        try:
            os.unlink(tmp_doc_path)
        except:
            pass


def _encontrar_inicio_secao_formulario(elementos_sem_sectpr):
    """Encontra onde a secao repetivel (paginas 2 e 3) comeca no modelo."""
    for i, elem in enumerate(elementos_sem_sectpr):
        texto = ''.join(elem.itertext())
        if '[DATA]' in texto:
            return max(0, i - 1)

    raise ValueError(
        'Nao foi possivel encontrar o inicio da secao de formulario no modelo2.docx. '
        'Verifique se o placeholder [DATA] existe nas paginas de formulario.'
    )


def _replace_in_element_text(elemento, old, new):
    """Substitui texto em todos os nos textuais do elemento XML."""
    for node in elemento.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        if node.text:
            node.text = node.text.replace(old, new)


def _inserir_antes_do_sectpr(body, elemento):
    for idx, child in enumerate(body):
        if child.tag.endswith('}sectPr'):
            body.insert(idx, elemento)
            return
    body.append(elemento)


def _criar_paragrafo_quebra_pagina():
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p = etree.Element(f'{{{ns_w}}}p')
    r = etree.SubElement(p, f'{{{ns_w}}}r')
    br = etree.SubElement(r, f'{{{ns_w}}}br')
    br.set(f'{{{ns_w}}}type', 'page')
    return p
